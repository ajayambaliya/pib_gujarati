import os
import aiohttp
import logging
from bs4 import BeautifulSoup
from deep_translator import GoogleTranslator
from pymongo import MongoClient
from docx import Document
from telegram import Bot
from telegram.error import TelegramError
import subprocess
import asyncio
import re
from io import BytesIO

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler()],
)

# Load environment variables
DB_NAME = os.getenv("DB_NAME")
COLLECTION_NAME = os.getenv("COLLECTION_NAME")
MONGO_CONNECTION_STRING = os.getenv("MONGO_CONNECTION_STRING")
TEMPLATE_URL = "https://docs.google.com/document/d/1GoHxD3FSM8-RhIJu_WGr4NVjVthCzpfx/export?format=docx"
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
TELEGRAM_CHANNEL_ID = os.getenv("TELEGRAM_CHANNEL_ID")
TELEGRAM_CHANNEL_URL = "https://t.me/pib_gujarati"

# Initialize MongoDB client
client = MongoClient(MONGO_CONNECTION_STRING)
db = client[DB_NAME]
collection = db[COLLECTION_NAME]

def chunk_text(text, max_length=4999):
    chunks = []
    current_chunk = ""
    for sentence in text.split(". "):
        if len(current_chunk) + len(sentence) + 2 <= max_length:
            current_chunk += sentence + ". "
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

async def download_template(url):
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                response.raise_for_status()
                template_bytes = BytesIO(await response.read())
                logging.info("Template downloaded successfully")
                return template_bytes
    except aiohttp.ClientError as e:
        logging.error(f"Error downloading template: {e}")
        return None

async def scrape_content():
    base_url = "https://pib.gov.in"
    main_url = f"{base_url}/allRel.aspx"

    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(main_url) as response:
                response.raise_for_status()
                soup = BeautifulSoup(await response.text(), "html.parser")

        links = []
        content_area = soup.find("div", class_="content-area")
        if content_area:
            for a_tag in content_area.find_all("a", href=True):
                href = a_tag["href"]
                full_link = f"{base_url}{href}"
                if not collection.find_one({"link": full_link}):
                    links.append(full_link)

        for link in links:
            async with aiohttp.ClientSession() as session:
                async with session.get(link) as response:
                    response.raise_for_status()
                    soup = BeautifulSoup(await response.text(), "html.parser")

            title_element = soup.find("h2")
            if title_element:
                title = title_element.get_text(strip=True)
            else:
                logging.warning(f"No <h2> found for link: {link}")
                continue

            content, content_gujarati = [], []
            asterisk_pattern = re.compile(r'\*{2,}')  # Regex pattern to match two or more asterisks

            for paragraph in soup.find_all("p", style="text-align:justify"):
                text = paragraph.get_text(strip=True)

                # Stop scraping if any asterisk pattern (e.g., **, ***) is found
                if asterisk_pattern.search(text):
                    logging.info(f"Stopping scraping due to asterisks in link: {link}")
                    break

                content.append(text)

                chunks = chunk_text(text)
                translated_chunks = []
                for chunk in chunks:
                    translated_chunk = GoogleTranslator(source="en", target="gu").translate(chunk)
                    if translated_chunk:
                        translated_chunks.append(translated_chunk)

                content_gujarati.append(" ".join(translated_chunks))

            if not content:
                for paragraph in soup.find_all("p", style="margin-left:0cm; margin-right:0cm; text-align:justify"):
                    text = paragraph.get_text(strip=True)
                    
                    # Stop scraping if any asterisk pattern (e.g., **, ***) is found
                    if asterisk_pattern.search(text):
                        logging.info(f"Stopping scraping due to asterisks in link: {link}")
                        break

                    content.append(text)

                    chunks = chunk_text(text)
                    translated_chunks = []
                    for chunk in chunks:
                        translated_chunk = GoogleTranslator(source="en", target="gu").translate(chunk)
                        if translated_chunk:
                            translated_chunks.append(translated_chunk)

                    content_gujarati.append(" ".join(translated_chunks))

            collection.insert_one({"link": link})

            if content:
                await generate_and_send_document(title, content, content_gujarati, link)
            else:
                await send_small_post_to_telegram(title, content, content_gujarati, link)

    except aiohttp.ClientError as e:
        logging.error(f"Error scraping content: {e}")

async def generate_and_send_document(title, content, content_gujarati, source_url):
    template_bytes = await download_template(TEMPLATE_URL)
    if not template_bytes:
        return

    try:
        doc = Document(template_bytes)
        doc.paragraphs[0].text = ""
        
        title_chunks = chunk_text(title)
        translated_title_chunks = []
        for chunk in title_chunks:
            translated_chunk = GoogleTranslator(source="en", target="gu").translate(chunk)
            if translated_chunk:
                translated_title_chunks.append(translated_chunk)
        translated_title = " ".join(translated_title_chunks)
        
        doc.add_heading(translated_title, 0)
        doc.paragraphs[1].text = title

        for eng_paragraph, guj_paragraph in zip(content, content_gujarati):
            if eng_paragraph.strip() and guj_paragraph.strip():
                para = doc.add_paragraph()
                para.add_run("• ").bold = True
                para.add_run(guj_paragraph)
                doc.add_paragraph("• " + eng_paragraph)
                doc.add_paragraph()

        promotional_message = "Don't miss out on the latest updates! Stay informed with our channel."
        doc.add_paragraph(promotional_message)
        doc.add_paragraph(f"Join our Telegram Channel for more updates: {TELEGRAM_CHANNEL_URL}")

        output_docx = "output.docx"
        doc.save(output_docx)

        pdf_file = await convert_docx_to_pdf(output_docx)
        pdf_name = f"{get_truncated_title(title)}.pdf"
        await send_to_telegram(pdf_file, pdf_name, f"🔖 {translated_title}\n\n🔗 Source: {source_url}\n\n{promotional_message}\n\n📥 Join our channel to get the latest updates: {TELEGRAM_CHANNEL_URL}")

    except Exception as e:
        logging.error(f"Error processing document: {e}")
    finally:
        cleanup_files(["output.docx", "output.pdf"])

async def send_small_post_to_telegram(title, content, content_gujarati, source_url):
    try:
        bot = Bot(token=TELEGRAM_BOT_TOKEN)
        
        title_chunks = chunk_text(title)
        translated_title_chunks = []
        for chunk in title_chunks:
            translated_chunk = GoogleTranslator(source="en", target="gu").translate(chunk)
            if translated_chunk:
                translated_title_chunks.append(translated_chunk)
        translated_title = " ".join(translated_title_chunks)
        
        message = f"🗞️ {translated_title}\n\n"
        for eng_paragraph, guj_paragraph in zip(content, content_gujarati):
            if eng_paragraph.strip() and guj_paragraph.strip():
                message += f"• {guj_paragraph}\n• {eng_paragraph}\n\n"

        promotional_message = "Don't miss out on the latest updates! Stay informed with our channel."
        message += f"🔗 Source: {source_url}\n\n{promotional_message}\n📥 Join our Telegram Channel for more updates: {TELEGRAM_CHANNEL_URL}"
        await bot.send_message(chat_id=TELEGRAM_CHANNEL_ID, text=message)

    except TelegramError as e:
        logging.error(f"Error sending small post to Telegram: {e}")
    except Exception as e:
        logging.error(f"Unexpected error sending small post to Telegram: {e}")

async def convert_docx_to_pdf(input_docx):
    output_pdf = "output.pdf"
    try:
        subprocess.run(["libreoffice", "--convert-to", "pdf", "--outdir", ".", input_docx], check=True)
        logging.info(f"Converted {input_docx} to PDF successfully")
        return output_pdf
    except subprocess.CalledProcessError as e:
        logging.error(f"Error converting DOCX to PDF: {e}")
        return None

def get_truncated_title(title):
    max_length = 200
    invalid_chars = r'\/:*?"<>|'
    sanitized_title = "".join([c for c in title if c not in invalid_chars])
    return sanitized_title[:max_length] if len(sanitized_title) > max_length else sanitized_title

async def send_to_telegram(pdf_file, pdf_name, caption):
    try:
        bot = Bot(token=TELEGRAM_BOT_TOKEN)
        with open(pdf_file, "rb") as file:
            await bot.send_document(chat_id=TELEGRAM_CHANNEL_ID, document=file, filename=pdf_name, caption=caption)
        logging.info(f"Sent {pdf_name} to Telegram successfully")
    except TelegramError as e:
        logging.error(f"Error sending document to Telegram: {e}")

def cleanup_files(files):
    for file in files:
        try:
            if os.path.exists(file):
                os.remove(file)
                logging.info(f"Removed file: {file}")
        except Exception as e:
            logging.error(f"Error removing file {file}: {e}")

if __name__ == "__main__":
    asyncio.run(scrape_content())
